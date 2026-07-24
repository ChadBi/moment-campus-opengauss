"""
生成"基于openGauss的此刻校园校园时空信息共享平台"课程设计报告
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION

def set_chinese_font(run, font_name='宋体', size=12, bold=False):
    """设置中文字体"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), font_name)

def add_title_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading('', level=level)
    run = heading.add_run(text)
    if level == 1:
        set_chinese_font(run, '黑体', 16, bold=True)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        set_chinese_font(run, '黑体', 14, bold=True)
    elif level == 3:
        set_chinese_font(run, '黑体', 12, bold=True)
    elif level == 4:
        set_chinese_font(run, '黑体', 12, bold=False)
    return heading

def add_paragraph(doc, text, indent=True):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_chinese_font(run, '宋体', 12)
    return p

def add_table(doc, headers, rows, table_id=''):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_chinese_font(run, '宋体', 10.5, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # 数据行
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_chinese_font(run, '宋体', 10.5)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    return table

def add_caption(doc, text, is_table=True):
    """添加图表标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_chinese_font(run, '宋体', 10.5)
    return p

def add_code_block(doc, code, language='sql'):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    return p

def setup_document(doc):
    """设置文档样式"""
    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)

def generate_report():
    doc = Document()
    setup_document(doc)
    
    # ==================== 封面 ====================
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('江 南 大 学')
    set_chinese_font(run, '黑体', 26, bold=True)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('课 程 设 计 报 告')
    set_chinese_font(run, '黑体', 22, bold=True)
    
    for _ in range(2):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('基于 openGauss 的"此刻校园"校园时空')
    set_chinese_font(run, '黑体', 18, bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('信息共享平台设计与实现')
    set_chinese_font(run, '黑体', 18, bold=True)
    
    for _ in range(4):
        doc.add_paragraph()
    
    info = [
        ('题   目：', '基于 openGauss 的"此刻校园"校园时空信息共享平台设计与实现'),
        ('院（系）：', '物联网工程学院'),
        ('专   业：', '计算机科学与技术'),
        ('班   级：', '计科2201班'),
        ('学   号：', '1042200000'),
        ('姓   名：', '张同学'),
        ('指导老师：', '李教授'),
    ]
    
    for label, value in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(6)
        run = p.add_run(label)
        set_chinese_font(run, '宋体', 14, bold=False)
        run = p.add_run(value)
        set_chinese_font(run, '宋体', 14)
    
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('课程设计时间：2026 年 6 月 — 2026 年 7 月')
    set_chinese_font(run, '宋体', 12)
    
    doc.add_page_break()
    
    # ==================== 分工说明 ====================
    add_title_heading(doc, '分工说明', level=2)
    add_paragraph(doc, '本项目为单人独立完成，涵盖需求分析、数据库设计（概念模型/逻辑模型/物理模型）、前后端开发、openGauss 物理模型落地与测试全流程。具体分工如下：')
    
    items = [
        '需求分析与数据字典编制：独立完成，包括用户调研、组织机构图绘制、数据流图设计、判定表与判定树设计、数据字典编写',
        '数据库概念模型设计：独立完成，包括21个实体识别、35个联系定义、6大功能模块划分、E-R图绘制',
        '数据库逻辑模型设计：独立完成，包括21张关系模式定义、3NF规范化验证、15个视图设计、完整性约束定义',
        '数据库物理模型设计：独立完成，包括4个表空间设计、66个索引优化、8个存储过程编写、8个触发器实现、4个物化视图、7张分区表',
        '后端开发：独立完成，基于FastAPI + SQLAlchemy 2.0异步框架，实现11个API模块、172项自动化测试',
        '前端开发：独立完成，基于React + TypeScript + Vite + Tailwind CSS，实现20+页面组件',
        'openGauss适配与部署：独立完成，包括SQLite到openGauss迁移、Docker部署、华为云混合部署验证',
        '文档编写：独立完成，包括全部设计文档与本课程设计报告',
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Pt(24)
        run = p.add_run('• ' + item)
        set_chinese_font(run, '宋体', 12)
    
    doc.add_page_break()
    
    # ==================== 目录 ====================
    add_title_heading(doc, '目  录', level=1)
    doc.add_paragraph()
    
    toc = [
        ('1  系统概述', '1'),
        ('2  数据库设计各阶段书面文档', '3'),
        ('3  开发环境与开发工具介绍', '6'),
        ('4  系统需求分析', '9'),
        ('  4.1  应用背景与调查方法', '9'),
        ('  4.2  用户痛点分析', '10'),
        ('  4.3  用户需求清单', '12'),
        ('  4.4  组织机构图', '14'),
        ('  4.5  数据流图', '16'),
        ('  4.6  判定表与判定树', '20'),
        ('  4.7  数据字典', '24'),
        ('5  功能需求分析与E-R图说明', '29'),
        ('  5.1  功能需求分析', '29'),
        ('  5.2  概念设计', '31'),
        ('  5.3  E-R图说明', '35'),
        ('6  系统设计', '40'),
        ('  6.1  模块设计', '40'),
        ('  6.2  逻辑结构设计', '45'),
        ('  6.3  物理结构设计', '52'),
        ('7  系统实现', '65'),
        ('  7.1  后端架构与核心接口', '65'),
        ('  7.2  前端界面实现', '72'),
        ('  7.3  数据库物理模型落地', '80'),
        ('8  系统运行情况说明与维护计划', '88'),
        ('  8.1  运行情况', '88'),
        ('  8.2  测试结果', '89'),
        ('  8.3  维护计划', '92'),
        ('9  小结', '95'),
        ('  9.1  开发收获', '95'),
        ('  9.2  遇到的问题及解决方法', '97'),
        ('10  参考文献', '100'),
        ('11  致谢', '102'),
        ('附录A  核心代码', '104'),
    ]
    
    for title, page in toc:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(title + ' ' + '.' * (60 - len(title) * 2) + ' ' + page)
        set_chinese_font(run, '宋体', 12)
    
    doc.add_page_break()
    
    return doc

if __name__ == '__main__':
    doc = generate_report()
    doc.save('d:/Project/database-class/moment-campus/docs/课程设计报告_part1.docx')
    print('Part 1 generated successfully')
